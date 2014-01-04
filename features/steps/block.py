# encoding: utf-8

"""
Step implementations for block content containers
"""

from behave import given, then, when

from docx import Document
from docx.table import Table

from .helpers import test_docx


# given ===================================================

@given('a document containing a table')
def given_a_document_containing_a_table(context):
    docx_path = test_docx('blk-containing-table')
    context.document = Document(docx_path)


@given('a paragraph')
def given_a_paragraph(context):
    context.document = Document()
    body = context.document.body
    context.p = body.add_paragraph()


# when ====================================================

@when('I add a paragraph')
def when_add_paragraph(context):
    body = context.document.body
    context.p = body.add_paragraph()


@when('I add a table')
def when_add_table(context):
    rows, cols = 2, 2
    context.document.body.add_table(rows, cols)


# then =====================================================

@then('I can access the table')
def then_can_access_table(context):
    table = context.document.body.tables[-1]
    assert isinstance(table, Table)


@then('the new table appears in the document')
def then_new_table_appears_in_document(context):
    table = context.document.body.tables[-1]
    assert isinstance(table, Table)
